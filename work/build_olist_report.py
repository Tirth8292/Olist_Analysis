from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\tirth\Documents\Codex\2026-06-23\hi-2\outputs\Olist_Analytics_Engineering_Project_Report.docx")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, header=True):
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if header and row_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.style.font.size = Pt(16)
        p.style.font.color.rgb = BLUE
    elif level == 2:
        p.style.font.size = Pt(13)
        p.style.font.color.rgb = BLUE
    else:
        p.style.font.size = Pt(12)
        p.style.font.color.rgb = DARK_BLUE
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    p.add_run(text)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [1.85, 4.65])
    table.rows[0].cells[0].text = "Area"
    table.rows[0].cells[1].text = "Details"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    style_table(table)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    p.add_run(f"\n{body}")
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("Olist Analytics Engineering Project Report")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE

    subtitle = doc.add_paragraph()
    subtitle.add_run("Project type: Modern data analytics platform for e-commerce marketplace intelligence").italic = True

    add_callout(
        doc,
        "Short summary",
        "This project converts 9 disconnected Olist e-commerce CSV files into a clean analytics system. "
        "It builds warehouse tables, business metrics, seller scoring, customer analysis, repeat-order insights, "
        "and a local API that can be tested through localhost.",
    )

    add_heading(doc, "1. Project Motive", 1)
    doc.add_paragraph(
        "The main motive of this project is to solve a common business problem: data exists in many separate files, "
        "but business teams need simple answers. Olist has data about customers, orders, sellers, products, payments, "
        "reviews, and delivery locations. Without a proper data model, teams must manually join files every time they "
        "need an answer."
    )
    doc.add_paragraph(
        "This project creates a repeatable pipeline that turns raw data into useful information for Growth, Operations, "
        "and Customer teams. It is designed like a real analytics engineering project, not only a notebook or one-time report."
    )

    add_heading(doc, "2. Data Used", 1)
    add_key_value_table(
        doc,
        [
            ("Dataset", "Brazilian E-Commerce Public Dataset by Olist"),
            ("Source files", "9 CSV files: customers, orders, order items, payments, reviews, products, sellers, geolocation, and product category translation"),
            ("Scale", "99,441 orders, 112,650 order items, 3,095 sellers, and about 1,000,000 geolocation rows"),
            ("Time period", "Historical e-commerce orders from 2016 to 2018"),
            ("Storage created", "Local SQLite warehouse: data/warehouse/olist.db"),
        ],
    )

    add_heading(doc, "3. What Was Built", 1)
    add_bullet(doc, "A raw ingestion layer that loads all CSV files into a local warehouse.")
    add_bullet(doc, "A staging layer that cleans column names, data types, dates, city names, and product categories.")
    add_bullet(doc, "A star schema with fact and dimension tables for orders, order items, payments, customers, sellers, products, dates, and geolocation.")
    add_bullet(doc, "Business views for category revenue, seller performance, and delivery SLA analysis.")
    add_bullet(doc, "Feature outputs for seller scores, customer RFM segments, cohort retention, delivery lanes, and repeat order patterns.")
    add_bullet(doc, "A FastAPI localhost API where users can test seller score, seller leaderboard, delivery estimate, and repeat-order endpoints.")

    add_heading(doc, "4. Project Architecture", 1)
    add_key_value_table(
        doc,
        [
            ("Step 1: Ingest", "Python reads the raw CSV files and loads them into the local database."),
            ("Step 2: Transform", "SQL creates clean staging tables and modeled fact/dimension tables."),
            ("Step 3: Analyze", "Python and SQL calculate seller scores, RFM, cohort retention, delivery lanes, and repeat orders."),
            ("Step 4: Serve", "FastAPI exposes selected insights through localhost endpoints."),
            ("Step 5: Present", "CSV reports and dashboard guidance can be used in Power BI."),
        ],
    )

    add_heading(doc, "5. Tools and Their Use", 1)
    add_key_value_table(
        doc,
        [
            ("Python", "Used for loading CSV files, feature engineering, quality checks, and API logic."),
            ("SQL", "Used for staging, fact tables, dimension tables, and business views."),
            ("SQLite", "Used as the local data warehouse so the project runs easily on one computer."),
            ("FastAPI", "Used to expose seller and delivery insights through localhost."),
            ("Apache Airflow", "Included as the orchestration layer to automate the pipeline steps in production."),
            ("Docker", "Included to make the project reproducible and easier to run in a clean environment."),
            ("dbt", "Included as the analytics engineering structure for production-style transformations and tests."),
            ("Power BI", "Planned for dashboard pages using generated CSV reports and warehouse tables."),
            ("Polars", "Not used in the current working version; it can be added later for faster large CSV processing."),
        ],
    )

    add_heading(doc, "6. Seller Performance Score", 1)
    doc.add_paragraph(
        "The seller leaderboard is based on a weighted performance score. This score helps identify strong sellers, "
        "watchlist sellers, and high-risk sellers before they hurt customer experience."
    )
    add_key_value_table(
        doc,
        [
            ("On-time delivery rate", "40% of the score. Sellers who deliver before or on the estimated date perform better."),
            ("Average review score", "30% of the score. Better customer reviews increase the seller score."),
            ("Order cancellation rate", "20% of the score. Lower cancellation improves the score."),
            ("Average order value growth", "10% of the score. Sellers with better revenue growth receive a small boost."),
        ],
    )

    add_heading(doc, "7. API Features on Localhost", 1)
    doc.add_paragraph(
        "The API runs locally at http://127.0.0.1:8000/docs. This means it works on the same computer where the project is running. "
        "The docs page is not a public website; it is a testing page for the local API."
    )
    add_key_value_table(
        doc,
        [
            ("GET /health", "Checks whether the API is running."),
            ("POST /seller/score", "Accepts a seller_id and returns score, risk tier, GMV, metrics, and recommendation."),
            ("GET /sellers/leaderboard", "Shows best or worst sellers based on the weighted performance score."),
            ("GET /seller/{seller_id}", "Shows detailed information for one seller, including rank and score breakdown."),
            ("GET /delivery/estimate", "Returns estimated delivery days for a historical origin and destination zip route."),
            ("GET /orders/repeat-by-city", "Shows which product categories are repeatedly ordered by customers in each city."),
        ],
    )

    add_heading(doc, "8. Important Files Used", 1)
    add_key_value_table(
        doc,
        [
            ("extract_load.py", "Loads all raw Olist CSV files into the local warehouse."),
            ("01_staging.sql", "Cleans and standardizes raw tables."),
            ("02_marts.sql", "Builds fact and dimension tables."),
            ("03_views.sql", "Creates business views for analytics."),
            ("feature_engineering.py", "Creates seller scores, RFM, cohort retention, and delivery lane reports."),
            ("api/main.py", "Runs the FastAPI localhost endpoints."),
            ("olist_pipeline_dag.py", "Shows how Airflow would automate the pipeline."),
            ("docker-compose.yml", "Shows how Docker can run project services."),
            ("reports folder", "Stores generated CSV outputs for dashboard and API use."),
        ],
    )

    add_heading(doc, "9. Business Value", 1)
    add_bullet(doc, "Operations teams can find seller regions and delivery routes with high late-delivery risk.")
    add_bullet(doc, "Growth teams can identify strong sellers and categories that deserve more marketplace visibility.")
    add_bullet(doc, "Customer teams can understand repeat buying behavior by city and product category.")
    add_bullet(doc, "Analysts can use modeled tables instead of manually joining many CSV files again and again.")
    add_bullet(doc, "The API proves that the project can serve insights, not only create static reports.")

    add_heading(doc, "10. Easy Interview Explanation", 1)
    doc.add_paragraph(
        "I built an end-to-end analytics engineering project using the Olist e-commerce dataset. "
        "I ingested 9 raw CSV files, cleaned and modeled them into fact and dimension tables, created SQL business views, "
        "engineered seller risk, customer RFM, cohort retention, delivery lane, and repeat-order features, and exposed key insights "
        "through a FastAPI localhost API. The project demonstrates data modeling, pipeline design, analytics, and dashboard readiness."
    )

    add_callout(
        doc,
        "Final project story",
        "This project turns raw marketplace data into decision-ready insights. It helps teams understand seller quality, delivery performance, "
        "customer behavior, repeat orders by city, and product category performance using a repeatable analytics pipeline.",
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Olist Analytics Engineering Project Report").font.size = Pt(9)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

